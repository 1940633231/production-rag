"""Word(.docx) 文档加载器：基于 python-docx 抽取纯文本。

依赖:
    pip install python-docx

特性:
    - 顺序提取所有段落文本，空行自动跳过
    - 表格内容按行展开，每行单元格用 \\t 分隔
    - .doc（旧二进制格式）不支持，需先转换为 .docx
"""
from pathlib import Path

from app.core.logger import get_logger
from app.ingestion.document import Document
from app.ingestion.loader.base import BaseLoader

logger = get_logger(__name__)


class WordLoader(BaseLoader):
    """Word(.docx) 文件加载器。"""

    def load(self, file_path) -> Document:
        path = Path(file_path)
        logger.info("加载 Word: %s", path)

        if path.suffix.lower() == ".doc":
            raise ValueError(
                "不支持 .doc 旧格式，请先转换为 .docx: {}".format(path)
            )

        try:
            import docx
        except ImportError as e:
            raise ImportError(
                "WordLoader 需要 python-docx: pip install python-docx"
            ) from e

        doc = docx.Document(str(path))
        parts = []

        # 段落
        para_count = 0
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
                para_count += 1

        # 表格内容
        table_count = len(doc.tables)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        content = "\n\n".join(parts)
        if not content.strip():
            logger.warning("Word 文档无可提取文本: %s", path)

        return Document(
            document_id=path.stem,
            content=content,
            metadata={
                "file_name": path.name,
                "file_path": str(path),
                "file_type": "docx",
                "paragraph_count": para_count,
                "table_count": table_count,
            },
        )
